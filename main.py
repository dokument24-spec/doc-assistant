import telebot
from telebot import types
import datetime
import os
import gspread
import time
import re
from oauth2client.service_account import ServiceAccountCredentials
from docx import Document
from flask import Flask, request

BOT_TOKEN = '7640880064:AAEOqKU4mWP06Ob96K3h4VDfrIhfK164Eg0'
ADMIN_ID = 5780051172
SHEET_NAME = "DocExpress_Заявки"

bot = telebot.TeleBot(BOT_TOKEN)
user_data = {}
current_step = {}
current_type = {}
last_submit_time = {}

# === Google Таблица ===
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds = ServiceAccountCredentials.from_json_keyfile_name("credentials.json", scope)
client = gspread.authorize(creds)
sheet = client.open(SHEET_NAME).sheet1

# === Сценарии ===
scenarios = {
    "Договор подряда": [
        "Как зовут заказчика?",
        "Как зовут исполнителя?",
        "Какие работы нужно выполнить?",
        "До какой даты нужно всё сделать?",
        "Сколько вы договорились за работу?",
        "Укажите номер телефона для связи"
    ],
    "Жалоба в УК": [
        "На каком адресе произошла проблема?",
        "Что произошло? (опишите своими словами)",
        "Когда это случилось?",
        "Есть ли доказательства? (фото/видео)?",
        "Номер телефона для связи"
    ],
    "Заявление в суд": [
        "Кто подаёт заявление (ваше ФИО)?",
        "Против кого подаёте заявление (ФИО)?",
        "Что произошло? (опишите суть ситуации простыми словами)",
        "Что вы хотите, чтобы суд решил?",
        "Ваш номер телефона"
    ],
    "Претензия продавцу": [
        "Где и что вы купили?",
        "Что случилось с товаром или услугой?",
        "Как вы хотите решить этот вопрос? (возврат, замена, ремонт)",
        "Контактный телефон"
    ],
    "Акт выполненных работ": [
        "Кто заказчик? (ФИО или организация)",
        "Кто исполнитель?",
        "Что именно было сделано?",
        "Дата завершения работ?",
        "Ваш номер телефона"
    ],
    "Доверенность": [
        "Кто доверяет? (ФИО)",
        "Кому доверяет? (ФИО)",
        "На какие действия даётся доверенность?",
        "На какой срок?",
        "Контактный номер"
    ],
    "Заявление о прописке": [
        "Кто подаёт заявление?",
        "Кого нужно прописать?",
        "Адрес прописки?",
        "Ваш номер телефона"
    ]
}

def is_valid_phone(phone):
    return re.match(r"^(\+7|8)\d{10}$", phone)

def can_submit(chat_id):
    now = time.time()
    if chat_id in last_submit_time and now - last_submit_time[chat_id] < 60:
        return False, int(60 - (now - last_submit_time[chat_id]))
    last_submit_time[chat_id] = now
    return True, 0

def generate_docx(doc_type, answers):
    doc = Document()
    doc.add_heading(doc_type, 0)
    for i, (q, a) in enumerate(answers, 1):
        doc.add_paragraph(f"{i}. {q}")
        doc.add_paragraph(f"   ➤ {a}", style='List Bullet')
    filename = f"{doc_type}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = f"/tmp/{filename}"
    doc.save(filepath)
    return filepath

@bot.message_handler(commands=['start'])
def start(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    for key in scenarios:
        markup.add(types.KeyboardButton(key))
    bot.send_message(message.chat.id, "👋 Добро пожаловать в «Оформлятор»!\nВыберите тип документа:", reply_markup=markup)

@bot.message_handler(func=lambda msg: msg.text in scenarios)
def start_scenario(message):
    cid = message.chat.id
    current_type[cid] = message.text
    current_step[cid] = 0
    user_data[cid] = []
    bot.send_message(cid, scenarios[message.text][0])

@bot.message_handler(func=lambda msg: msg.chat.id in current_step)
def handle_response(message):
    cid = message.chat.id
    step = current_step[cid]
    scenario = scenarios[current_type[cid]]
    user_data[cid].append((scenario[step], message.text))

    if "телефон" in scenario[step].lower():
        if not is_valid_phone(message.text):
            bot.send_message(cid, "❗️Введите номер телефона в формате +7XXXXXXXXXX или 8XXXXXXXXXX")
            return

    step += 1
    if step < len(scenario):
        current_step[cid] = step
        bot.send_message(cid, scenario[step])
    else:
        bot.send_message(cid, "📋 Проверьте ваши ответы:")
        for q, a in user_data[cid]:
            bot.send_message(cid, f"{q}\n➤ {a}")
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
        markup.add("✅ Да", "❌ Нет")
        bot.send_message(cid, "Отправить этот документ?", reply_markup=markup)
        current_step.pop(cid)

@bot.message_handler(func=lambda msg: msg.text in ["✅ Да", "❌ Нет"])
def confirm_send(message):
    cid = message.chat.id
    if message.text == "✅ Да":
        allowed, wait = can_submit(cid)
        if not allowed:
            bot.send_message(cid, f"⏱ Подождите {wait} секунд перед новой заявкой.")
            return

        doc_type = current_type[cid]
        answers = user_data[cid]
        filepath = generate_docx(doc_type, answers)

        with open(filepath, "rb") as f:
            bot.send_document(cid, f)

        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        try:
            sheet.append_row([now, cid, doc_type] + [a for _, a in answers])
        except Exception as e:
            bot.send_message(ADMIN_ID, f"⚠️ Ошибка записи в таблицу: {e}")

        user = message.from_user
        summary = f"📥 Новая заявка от @{user.username or 'Без ника'}\nТип документа: {doc_type}\n"
        for q, a in answers:
            summary += f"\n{q}\n➤ {a}"
        bot.send_message(ADMIN_ID, summary)
        bot.send_message(cid, "✅ Документ создан и отправлен. Спасибо!")
    else:
        bot.send_message(cid, "🔁 Хорошо, начнём заново.")
        start(message)

@bot.message_handler(content_types=['voice'])
def handle_voice(message):
    cid = message.chat.id
    bot.send_message(cid, "🎤 Голосовые пока не обрабатываю, но скоро научусь. Напишите текстом 🙏")
    bot.send_voice(ADMIN_ID, message.voice.file_id, caption=f"🎧 Голосовое от {cid}")

@bot.message_handler(content_types=['photo', 'document', 'audio', 'video'])
def handle_files(message):
    cid = message.chat.id
    doc_type = current_type.get(cid, "Без сценария")
    username = f"@{message.from_user.username}" if message.from_user.username else "без ника"

    file_info = None
    if message.content_type == 'photo':
        file_info = bot.get_file(message.photo[-1].file_id)
        file_type = "📷 Фото"
    elif message.content_type == 'document':
        file_info = bot.get_file(message.document.file_id)
        file_type = f"📄 Документ: {message.document.file_name}"
    elif message.content_type == 'audio':
        file_info = bot.get_file(message.audio.file_id)
        file_type = f"🎵 Аудио: {message.audio.title or 'без названия'}"
    elif message.content_type == 'video':
        file_info = bot.get_file(message.video.file_id)
        file_type = "🎞 Видео"

    if file_info:
        downloaded_file = bot.download_file(file_info.file_path)
        filename = file_info.file_path.split('/')[-1]
        filepath = f"/tmp/{filename}"
        with open(filepath, 'wb') as f:
            f.write(downloaded_file)

        bot.send_message(cid, f"{file_type} получен ✅")
        with open(filepath, 'rb') as f:
            caption = f"📎 Файл от {username}\nID: {cid}\nТип документа: {doc_type}"
            bot.send_document(ADMIN_ID, f, caption=caption)

# === Flask Webhook ===
app = Flask(__name__)

@app.route('/')
def index():
    return "Оформлятор работает."

@app.route(f"/{BOT_TOKEN}", methods=["POST"])
def webhook():
    update = telebot.types.Update.de_json(request.stream.read().decode("utf-8"))
    bot.process_new_updates([update])
    return "ok", 200

if __name__ == '__main__':
    app.run(debug=True)

